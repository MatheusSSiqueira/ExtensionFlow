# -*- coding: utf-8 -*-
"""
ExtensionFlow
=============
Sistema RAG multi-agente para auxiliar alunos no processo de envio de
atividades extensionistas, com base em documentos institucionais (DOCX/PDF).

Agentes:
    - ManualAnalystAgent : recupera chunks via FAISS e extrai regras com LLM.
    - AttendantAgent     : gerencia a interação com o aluno e entrega respostas.

Provedores de LLM suportados:
    "openai"    → ChatGPT / GPT-4o          (OPENAI_API_KEY)
    "anthropic" → Claude Sonnet             (ANTHROPIC_API_KEY)
    "gemini"    → Google Gemini 1.5 Pro     (GOOGLE_API_KEY)

Dependências:
    pip install langchain langchain-text-splitters langchain-core
                langchain-openai langchain-anthropic
                langchain-google-genai langchain-community
                faiss-cpu pdfplumber python-docx tiktoken python-dotenv

Uso:
    python extensionflow.py
"""

# ---------------------------------------------------------------------------
# Imports da biblioteca padrão
# ---------------------------------------------------------------------------
import logging
import os
import re
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Dependências opcionais de terceiros (verificadas em tempo de execução)
# ---------------------------------------------------------------------------
try:
    import pdfplumber
except ImportError as exc:  # pragma: no cover
    raise ImportError("Instale pdfplumber: pip install pdfplumber") from exc

try:
    from docx import Document as DocxDocument
except ImportError as exc:  # pragma: no cover
    raise ImportError("Instale python-docx: pip install python-docx") from exc

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv é opcional; variáveis de ambiente são suficientes

from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ---------------------------------------------------------------------------
# Configurações globais
# ---------------------------------------------------------------------------
_ALLOWED_PROVIDERS = frozenset({"openai", "anthropic", "gemini"})
_MAX_QUERY_LENGTH  = 2000    # caracteres — prevenção de prompt injection
_API_KEY_PATTERNS  = {
    "openai":    re.compile(r"^sk-[A-Za-z0-9\-_]{20,}$"),
    "anthropic": re.compile(r"^sk-ant-[A-Za-z0-9\-_]{20,}$"),
    "gemini":    re.compile(r"^AIza[A-Za-z0-9\-_]{35,}$"),
}

CONTENT_DIR      = Path(os.getenv("CONTENT_DIR",      "/content")).resolve()
MEMORY_FILE      = CONTENT_DIR / "regras_memoria.txt"
FAISS_INDEX_PATH = Path(os.getenv("FAISS_INDEX_PATH", "/content/faiss_index")).resolve()
LLM_PROVIDER     = os.getenv("LLM_PROVIDER", "openai").lower().strip()

CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 200
RETRIEVER_K   = 4

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Utilitários de segurança
# ---------------------------------------------------------------------------

def _validate_provider(provider: str) -> str:
    """Valida o provedor contra lista de permissão (allowlist)."""
    p = provider.lower().strip()
    if p not in _ALLOWED_PROVIDERS:
        raise ValueError(
            f"Provedor '{provider}' não suportado. "
            f"Use: {', '.join(sorted(_ALLOWED_PROVIDERS))}."
        )
    return p


def _validate_api_key(provider: str, key: str) -> str:
    """
    Valida o formato da chave de API contra o padrão esperado.
    Nunca registra a chave em logs.
    """
    if not key:
        raise EnvironmentError(
            f"Chave de API não encontrada para o provedor '{provider}'.\n"
            f"Defina a variável de ambiente correspondente ou use um arquivo .env."
        )
    pattern = _API_KEY_PATTERNS.get(provider)
    if pattern and not pattern.match(key):
        # Exibe apenas os 6 primeiros caracteres para diagnóstico
        preview = key[:6] + "..." if len(key) > 6 else "***"
        log.warning(
            "Chave '%s' pode estar em formato inválido para '%s'. "
            "Verifique se foi copiada corretamente.",
            preview,
            provider,
        )
    return key


def _sanitize_query(query: str) -> str:
    """
    Sanitiza a consulta do aluno:
      - Limita o comprimento para evitar prompt injection.
      - Remove caracteres de controle.
    """
    sanitized = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", query)
    if len(sanitized) > _MAX_QUERY_LENGTH:
        log.warning(
            "Consulta truncada de %d para %d caracteres.",
            len(sanitized),
            _MAX_QUERY_LENGTH,
        )
        sanitized = sanitized[:_MAX_QUERY_LENGTH]
    return sanitized.strip()


def _safe_path(target: Path, base: Path) -> Path:
    """
    Garante que `target` está dentro de `base` (prevenção de path traversal).
    Levanta ValueError se o caminho escapar do diretório base.
    """
    resolved = target.resolve()
    if not str(resolved).startswith(str(base.resolve())):
        raise ValueError(
            f"Acesso negado: '{target}' está fora do diretório permitido '{base}'."
        )
    return resolved


# ---------------------------------------------------------------------------
# Fábrica de LLM — OpenAI (ChatGPT) | Anthropic (Claude) | Google (Gemini)
# ---------------------------------------------------------------------------

def build_llm(provider: str = LLM_PROVIDER, api_key: str = ""):
    """
    Instancia o LLM conforme o provedor escolhido.

    As importações são feitas sob demanda (lazy) para que apenas as
    dependências do provedor escolhido precisem estar instaladas.

    Args:
        provider : "openai", "anthropic" ou "gemini".
        api_key  : Chave de API. Se omitida, lida da variável de ambiente.

    Retorna:
        Instância do chat model LangChain correspondente.
    """
    provider = _validate_provider(provider)

    if provider == "openai":
        # pylint: disable=import-outside-toplevel
        from langchain_openai import ChatOpenAI
        key = _validate_api_key(provider, api_key or os.getenv("OPENAI_API_KEY", ""))
        log.info("LLM: OpenAI GPT-4o inicializado.")
        return ChatOpenAI(model="gpt-4o", temperature=0.2, api_key=key)

    if provider == "anthropic":
        # pylint: disable=import-outside-toplevel
        from langchain_anthropic import ChatAnthropic
        key = _validate_api_key(provider, api_key or os.getenv("ANTHROPIC_API_KEY", ""))
        log.info("LLM: Anthropic Claude Sonnet inicializado.")
        return ChatAnthropic(model="claude-sonnet-4-5", temperature=0.2, api_key=key)

    # provider == "gemini"
    # pylint: disable=import-outside-toplevel
    from langchain_google_genai import ChatGoogleGenerativeAI
    key = _validate_api_key(provider, api_key or os.getenv("GOOGLE_API_KEY", ""))
    log.info("LLM: Google Gemini 1.5 Pro inicializado.")
    return ChatGoogleGenerativeAI(
        model="gemini-1.5-pro",
        temperature=0.2,
        google_api_key=key,
    )


def build_embeddings(provider: str = LLM_PROVIDER, api_key: str = ""):
    """
    Instancia o modelo de embeddings adequado ao provedor.

      openai / anthropic → OpenAI  text-embedding-3-small
      gemini             → Google  models/embedding-001 (sem custo extra)

    Args:
        provider : "openai", "anthropic" ou "gemini".
        api_key  : Chave de API. Se omitida, lida da variável de ambiente.
    """
    provider = _validate_provider(provider)

    if provider == "gemini":
        # pylint: disable=import-outside-toplevel
        from langchain_google_genai import GoogleGenerativeAIEmbeddings
        key = _validate_api_key(provider, api_key or os.getenv("GOOGLE_API_KEY", ""))
        log.info("Embeddings: Google models/embedding-001.")
        return GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=key,
        )

    # openai ou anthropic → embeddings OpenAI
    # pylint: disable=import-outside-toplevel
    from langchain_openai import OpenAIEmbeddings
    key = _validate_api_key("openai", api_key or os.getenv("OPENAI_API_KEY", ""))
    log.info("Embeddings: OpenAI text-embedding-3-small.")
    return OpenAIEmbeddings(model="text-embedding-3-small", api_key=key)


# ---------------------------------------------------------------------------
# Extração de texto — DOCX e PDF
# ---------------------------------------------------------------------------

def extract_text_from_docx(docx_path: Path) -> Optional[str]:
    """Extrai o texto de um arquivo DOCX parágrafo a parágrafo."""
    try:
        doc = DocxDocument(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs) if paragraphs else None
    except (OSError, ValueError, KeyError) as exc:
        log.error("Erro ao ler DOCX '%s': %s", docx_path.name, exc)
        return None


def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """Extrai o texto de todas as páginas de um arquivo PDF."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pages = [pg.extract_text() for pg in pdf.pages if pg.extract_text()]
        return "\n".join(pages) if pages else None
    except (OSError, ValueError) as exc:
        log.error("Erro ao ler PDF '%s': %s", pdf_path.name, exc)
        return None


# ---------------------------------------------------------------------------
# Pipeline de documentos
# ---------------------------------------------------------------------------

def load_documents(content_dir: Path) -> dict[str, str]:
    """
    Escaneia `content_dir` por arquivos DOCX e PDF e retorna
    {nome_arquivo: conteudo_texto}.

    Apenas arquivos dentro de `content_dir` são processados
    (proteção contra path traversal em nomes de arquivo).
    """
    if not content_dir.exists():
        log.warning("Diretório '%s' não encontrado.", content_dir)
        return {}

    extractors: dict[str, tuple] = {
        ".docx": (extract_text_from_docx, "DOCX"),
        ".pdf":  (extract_text_from_pdf,  "PDF"),
    }
    result: dict[str, str] = {}

    for file in sorted(content_dir.iterdir()):
        if file.suffix not in extractors:
            continue
        try:
            _safe_path(file, content_dir)      # bloqueia path traversal
        except ValueError as exc:
            log.warning("Arquivo ignorado: %s", exc)
            continue

        fn, label = extractors[file.suffix]
        content = fn(file)
        if content:
            result[file.name] = content
            log.info("[%s] '%s' — %d chars.", label, file.name, len(content))
        else:
            log.warning("[%s] '%s' vazio ou ilegível.", label, file.name)

    log.info("Documentos carregados: %d.", len(result))
    return result


def build_chunks(documents: dict[str, str]) -> list[Document]:
    """Converte o dict de documentos em chunks LangChain."""
    langchain_docs = [
        Document(page_content=content, metadata={"source": name})
        for name, content in documents.items()
    ]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        add_start_index=True,
    )
    chunk_list = splitter.split_documents(langchain_docs)
    log.info("%d documento(s) → %d chunk(s).", len(langchain_docs), len(chunk_list))
    return chunk_list


# ---------------------------------------------------------------------------
# Banco vetorial FAISS
# ---------------------------------------------------------------------------

def build_vectorstore(chunk_list: list[Document], index_path: Path, embeddings):
    """
    Cria o índice FAISS a partir dos chunks e salva em disco.
    Se o índice já existir, é carregado sem reprocessar.

    Nota de segurança: allow_dangerous_deserialization=True é necessário
    porque o FAISS usa pickle internamente. O risco é aceitável aqui pois
    o índice é criado e consumido pela mesma aplicação no mesmo ambiente.
    Nunca carregue índices FAISS de fontes não confiáveis.
    """
    # pylint: disable=import-outside-toplevel
    from langchain_community.vectorstores import FAISS

    if index_path.exists():
        log.info("Carregando índice FAISS de '%s'...", index_path)
        vector_store = FAISS.load_local(
            str(index_path),
            embeddings,
            allow_dangerous_deserialization=True,  # seguro: índice gerado localmente
        )
        log.info("Índice FAISS carregado.")
        return vector_store

    log.info("Criando índice FAISS (pode levar alguns segundos)...")
    vector_store = FAISS.from_documents(chunk_list, embeddings)
    index_path.mkdir(parents=True, exist_ok=True)
    vector_store.save_local(str(index_path))
    log.info("Índice FAISS salvo em '%s'.", index_path)
    return vector_store


# ---------------------------------------------------------------------------
# Memória do agente
# ---------------------------------------------------------------------------

def load_memory_rules(memory_file: Path) -> str:
    """Lê o arquivo de memória; cria-o se não existir."""
    if not memory_file.exists():
        memory_file.touch()
        log.info("Arquivo de memória criado: '%s'.", memory_file)
        return ""
    content = memory_file.read_text(encoding="utf-8").strip()
    log.info("Memória '%s' %s.", memory_file.name, "carregada" if content else "vazia")
    return content


def append_to_memory(rule: str, memory_file: Path = MEMORY_FILE) -> None:
    """Persiste uma nova regra aprendida no arquivo de memória."""
    sanitized_rule = _sanitize_query(rule)
    if not sanitized_rule:
        log.warning("Regra vazia ignorada.")
        return
    with open(memory_file, "a", encoding="utf-8") as fh:
        fh.write(sanitized_rule + "\n")
    log.info("Regra registrada na memória.")


# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

ANALYST_PROMPT = ChatPromptTemplate.from_messages([
    (
        "system",
        """Você é o Agente Analista de Manuais do ExtensionFlow.
Sua ÚNICA fonte de informação são os trechos dos documentos institucionais abaixo.

REGRAS ABSOLUTAS:
- Nunca invente regras, prazos ou requisitos ausentes nos trechos fornecidos.
- Não realize o trabalho acadêmico do aluno.
- Se a informação não constar nos trechos, declare: "Informação não encontrada."
- Cite o nome do arquivo de origem de cada informação extraída.

REGRAS APRENDIDAS (memória):
{memory_rules}

TRECHOS RECUPERADOS:
{context}""",
    ),
    (
        "human",
        "Curso: {curso}\nConsulta: {question}\n\n"
        "Extraia e estruture em tópicos todas as regras, prazos e requisitos relevantes.",
    ),
])

ATTENDANT_PROMPT = ChatPromptTemplate.from_messages([
    (
        "system",
        """Você é o Agente Atendente do ExtensionFlow, sistema de suporte a alunos.
Transforme o resumo técnico do Agente Analista em uma resposta clara e amigável.

REGRAS:
- Use linguagem simples e encorajadora.
- Organize em passos numerados ou tópicos com marcadores (•).
- Não adicione informações além do que está no resumo do analista.
- Se o analista não encontrou informação, oriente o aluno a contatar a coordenação.
- Inclua o template de lista de presença APENAS se a atividade for presencial.""",
    ),
    (
        "human",
        "Curso: {curso}\nConsulta: {question}\n\n"
        "Resumo do Analista:\n{analyst_summary}\n\n"
        "Elabore a resposta final para o aluno.",
    ),
])


# ---------------------------------------------------------------------------
# ManualAnalystAgent
# ---------------------------------------------------------------------------

class ManualAnalystAgent:
    """
    Agente Analista de Manuais do ExtensionFlow.

    Usa busca semântica FAISS para recuperar os chunks mais relevantes
    e um LLM para extrair e estruturar regras com guardrails no prompt.
    """

    def __init__(self, retriever, llm, memory_rules: str = "") -> None:
        self._retriever   = retriever
        self._memory      = memory_rules or "Nenhuma regra aprendida ainda."

        self._chain = (
            {
                "context":      lambda inp: self._format_docs(
                                    self._retriever.invoke(inp["question"])
                                ),
                "question":     RunnablePassthrough() | (lambda inp: inp["question"]),
                "curso":        RunnablePassthrough() | (lambda inp: inp["curso"]),
                "memory_rules": lambda _: self._memory,
            }
            | ANALYST_PROMPT
            | llm
            | StrOutputParser()
        )
        log.info("ManualAnalystAgent pronto.")

    @staticmethod
    def _format_docs(docs: list[Document]) -> str:
        if not docs:
            return "Nenhum trecho relevante encontrado."
        return "\n\n".join(
            f"[Trecho {i} — {doc.metadata.get('source', '?')}]\n{doc.page_content}"
            for i, doc in enumerate(docs, 1)
        )

    def analyze(self, student_query: str, curso: str) -> dict:
        """
        Executa busca semântica + extração via LLM.

        Args:
            student_query : Consulta sanitizada do aluno.
            curso         : Nome do curso para contextualizar o prompt.

        Retorna:
            dict com 'status' e 'summary' (success) ou 'message' (fallback).
        """
        query = _sanitize_query(student_query)
        log.info("ManualAnalystAgent: analisando consulta.")
        try:
            summary = self._chain.invoke({"question": query, "curso": curso})
            return {"status": "success", "summary": summary}
        except (ValueError, RuntimeError, OSError) as exc:
            log.error("Erro no ManualAnalystAgent: %s", exc)
            return {"status": "fallback", "message": str(exc)}

    def update_memory(self, new_rules: str) -> None:
        """
        Atualiza as regras de memória do agente em tempo de execução.
        Útil após o AttendantAgent persistir uma nova correção via log_correction().

        Args:
            new_rules: Conteúdo atualizado do arquivo regras_memoria.txt.
        """
        if not new_rules.strip():
            return
        self._memory = new_rules.strip()
        log.info("ManualAnalystAgent: memória atualizada (%d chars).", len(self._memory))


# ---------------------------------------------------------------------------
# AttendantAgent
# ---------------------------------------------------------------------------

class AttendantAgent:
    """
    Agente Atendente do ExtensionFlow.

    Orquestra o ciclo ReAct completo:
        Observe  → valida contexto e sanitiza entrada
        Reason   → delega ao ManualAnalystAgent e interpreta o resultado
        Act      → formata via LLM e entrega resposta ao aluno
    """

    _REQUIRED_FIELDS = ("curso",)

    def __init__(self, analyst: ManualAnalystAgent, llm, memory_file: Path = MEMORY_FILE) -> None:
        self._analyst     = analyst
        self._memory_file = memory_file
        self._chain       = ATTENDANT_PROMPT | llm | StrOutputParser()
        log.info("AttendantAgent pronto.")

    def _validate_context(self, ctx: dict) -> Optional[str]:
        for field in self._REQUIRED_FIELDS:
            if not ctx.get(field, "").strip():
                return (
                    "Para te ajudar melhor, preciso saber: **qual é o seu curso?**\n"
                    "(Ex.: Análise e Desenvolvimento de Sistemas)"
                )
        return None

    def handle(self, student_query: str, context: Optional[dict] = None) -> str:
        """
        Ponto de entrada principal — ciclo ReAct completo.

        Args:
            student_query : Pergunta do aluno.
            context       : {"curso": "..."}.

        Retorna:
            Resposta estruturada para o aluno.
        """
        ctx = context or {}

        # Observe: validar e sanitizar
        missing = self._validate_context(ctx)
        if missing:
            return missing

        query  = _sanitize_query(student_query)
        curso  = _sanitize_query(ctx["curso"])

        # Act: delegar ao ManualAnalystAgent
        result = self._analyst.analyze(query, curso)

        if result["status"] == "fallback":
            return (
                "Não consegui recuperar informações dos manuais.\n\n"
                f"**Detalhe:** {result['message']}\n\n"
                "Tente reformular a pergunta ou entre em contato com a coordenação."
            )

        # Reason + Act: refinar via LLM
        try:
            return self._chain.invoke({
                "curso":           curso,
                "question":        query,
                "analyst_summary": result["summary"],
            })
        except (ValueError, RuntimeError, OSError) as exc:
            log.error("Erro no AttendantAgent ao formatar resposta: %s", exc)
            return (
                "Erro ao formatar a resposta. "
                "Resumo bruto do analista:\n\n" + result["summary"]
            )

    def log_correction(self, rule: str) -> None:
        """Persiste uma correção na memória de longo prazo."""
        append_to_memory(rule, self._memory_file)


# ---------------------------------------------------------------------------
# Execução principal (sem interface Gradio)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    _DEMO_QUERY   = "Qual é o prazo e o formato para envio da atividade de extensão?"
    _DEMO_CONTEXT = {"curso": "Análise e Desenvolvimento de Sistemas"}

    _docs = load_documents(CONTENT_DIR)
    if not _docs:
        raise SystemExit(f"Nenhum documento encontrado em '{CONTENT_DIR}'.")

    _chunks     = build_chunks(_docs)
    _embeddings = build_embeddings(LLM_PROVIDER)
    _vs         = build_vectorstore(_chunks, FAISS_INDEX_PATH, _embeddings)
    _retriever  = _vs.as_retriever(search_kwargs={"k": RETRIEVER_K})
    _llm        = build_llm(LLM_PROVIDER)
    _memory     = load_memory_rules(MEMORY_FILE)

    _analyst   = ManualAnalystAgent(retriever=_retriever, llm=_llm, memory_rules=_memory)
    _attendant = AttendantAgent(analyst=_analyst, llm=_llm)

    print("\n" + "=" * 65)
    print(f"CONSULTA : {_DEMO_QUERY}")
    print(f"CURSO    : {_DEMO_CONTEXT['curso']}")
    print("=" * 65)
    print(_attendant.handle(student_query=_DEMO_QUERY, context=_DEMO_CONTEXT))
    print("=" * 65)
